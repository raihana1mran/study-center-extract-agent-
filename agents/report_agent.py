"""
agents/report_agent.py — Generates DOCX, XLSX, and PDF reports.

Document structure:
  STATE
   └─ DISTRICT
        ├─ AI Code
        ├─ Study Centre Name
        ├─ Address

Uses: Gemma 4 31B (google/gemma-4-31b-it:free) for executive summary generation.
Primary report generation is deterministic (python-docx, openpyxl, fpdf2).
"""

from __future__ import annotations

import io
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any

from config import (
    DOCX_PATH, XLSX_PATH, PDF_PATH,
    MODELS, OPENROUTER_API_KEY, OPENROUTER_BASE_URL,
)
from utils.logger import log


# ─── OpenRouter client ──────────────────────────────────────────

def _get_ai_client():
    from openai import OpenAI
    return OpenAI(
        api_key=OPENROUTER_API_KEY,
        base_url=OPENROUTER_BASE_URL,
        default_headers={
            "HTTP-Referer": "https://github.com/nios-agent",
            "X-Title": "NIOS Study Centre Agent",
        },
    )


class ReportAgent:
    """
    Generates three report formats from DB study centre data.
    All reports follow the STATE → DISTRICT → Centre hierarchy.
    """

    def generate_all(self, centres: List[Any]) -> Dict[str, Path]:
        """
        Generate DOCX, XLSX, and PDF from a list of StudyCentre ORM objects.
        Returns: dict with keys 'docx', 'xlsx', 'pdf' pointing to output paths.
        """
        log.info(f"Generating reports for {len(centres)} study centres...")

        # Group by state → district
        grouped = self._group(centres)

        results = {}
        try:
            self._generate_docx(grouped, centres)
            results["docx"] = DOCX_PATH
            log.info(f"  ✓ DOCX: {DOCX_PATH}")
        except Exception as e:
            log.error(f"  ✗ DOCX generation failed: {e}")

        try:
            self._generate_xlsx(grouped, centres)
            results["xlsx"] = XLSX_PATH
            log.info(f"  ✓ XLSX: {XLSX_PATH}")
        except Exception as e:
            log.error(f"  ✗ XLSX generation failed: {e}")

        try:
            self._generate_pdf(grouped, centres)
            results["pdf"] = PDF_PATH
            log.info(f"  ✓ PDF:  {PDF_PATH}")
        except Exception as e:
            log.error(f"  ✗ PDF generation failed: {e}")

        return results

    # ── Grouping ───────────────────────────────────────────────

    def _group(self, centres: List[Any]) -> Dict[str, Dict[str, List]]:
        """Group centres into state → district → [centres]."""
        grouped = defaultdict(lambda: defaultdict(list))
        for c in centres:
            state    = (c.state    or "Unknown State").strip()
            district = (c.district or "Unknown District").strip()
            grouped[state][district].append(c)
        return {
            state: dict(districts)
            for state, districts in sorted(grouped.items())
        }

    # ── DOCX Report ───────────────────────────────────────────

    def _generate_docx(self, grouped: Dict, all_centres: List) -> None:
        """Generate a formatted Word document with STATE → DISTRICT hierarchy."""
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # ── Cover page ──
        doc.add_heading("NIOS Academic Study Centre Directory", level=0)
        doc.add_heading("Complete India-Wide Directory", level=1)

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(
            f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}\n"
            f"Total States: {len(grouped)} | "
            f"Total Centres: {len(all_centres)}\n"
            f"Category: Academic | Source: sdmis.nios.ac.in"
        )
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        doc.add_page_break()

        # ── Table of Contents placeholder ──
        doc.add_heading("Contents", level=1)
        for state in sorted(grouped.keys()):
            district_count  = len(grouped[state])
            centre_count    = sum(len(v) for v in grouped[state].values())
            p = doc.add_paragraph(
                f"  {state}  -  {district_count} districts, {centre_count} centres",
                style="List Bullet",
            )

        doc.add_page_break()

        # ── Per-State Sections ──
        for state, districts in sorted(grouped.items()):
            # State heading
            doc.add_heading(f"📍 {state}", level=1)
            total_in_state = sum(len(v) for v in districts.values())
            doc.add_paragraph(
                f"Districts: {len(districts)}  |  Total Academic Study Centres: {total_in_state}"
            ).runs[0].font.italic = True

            for district, centres_list in sorted(districts.items()):
                # District heading
                doc.add_heading(f"  {district}  ({len(centres_list)} centres)", level=2)

                # Table: AI Code | Name | Address
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"

                # Header row
                hdr = table.rows[0].cells
                hdr[0].text = "AI Code"
                hdr[1].text = "Study Centre Name"
                hdr[2].text = "Address"
                for cell in hdr:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Data rows
                for centre in sorted(centres_list, key=lambda x: x.ai_code):
                    row = table.add_row().cells
                    row[0].text = centre.ai_code    or ""
                    row[1].text = centre.name       or ""
                    row[2].text = centre.address    or ""

                # Column widths
                for row in table.rows:
                    row.cells[0].width = Inches(1.0)
                    row.cells[1].width = Inches(2.5)
                    row.cells[2].width = Inches(3.5)

                doc.add_paragraph("")  # spacing

            doc.add_page_break()

        doc.save(str(DOCX_PATH))

    # ── XLSX Report ───────────────────────────────────────────

    def _generate_xlsx(self, grouped: Dict, all_centres: List) -> None:
        """Generate a multi-sheet Excel workbook. One sheet per state + a Summary sheet."""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        wb.remove(wb.active)  # remove default sheet

        # ── Style definitions ──
        header_font   = Font(bold=True, color="FFFFFF", size=11)
        header_fill   = PatternFill("solid", fgColor="1B4F72")
        state_font    = Font(bold=True, color="FFFFFF", size=10)
        state_fill    = PatternFill("solid", fgColor="2980B9")
        district_font = Font(bold=True, size=10)
        district_fill = PatternFill("solid", fgColor="D6EAF8")
        center_align  = Alignment(horizontal="center", vertical="center", wrap_text=True)
        left_align    = Alignment(horizontal="left", vertical="top", wrap_text=True)
        thin_border   = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        )

        def style_cell(cell, font=None, fill=None, align=None, border=None):
            if font:   cell.font      = font
            if fill:   cell.fill      = fill
            if align:  cell.alignment = align
            if border: cell.border    = border

        # ── Summary sheet ──
        ws_summary = wb.create_sheet("Summary", 0)
        ws_summary.column_dimensions["A"].width = 35
        ws_summary.column_dimensions["B"].width = 15
        ws_summary.column_dimensions["C"].width = 15

        ws_summary.append(["NIOS Academic Study Centre Directory"])
        ws_summary["A1"].font = Font(bold=True, size=14)
        ws_summary.append([f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}"])
        ws_summary.append([f"Total Centres: {len(all_centres)}"])
        ws_summary.append([])
        ws_summary.append(["State / UT", "Districts", "Centres"])
        for cell in ws_summary["A5:C5"][0]:
            style_cell(cell, font=header_font, fill=header_fill, align=center_align)

        for state, districts in sorted(grouped.items()):
            cnt = sum(len(v) for v in districts.values())
            ws_summary.append([state, len(districts), cnt])

        # ── Per-state sheets ──
        for state, districts in sorted(grouped.items()):
            # Truncate sheet name to 31 chars (Excel limit)
            sheet_name = state[:31]
            ws = wb.create_sheet(sheet_name)

            # State header
            ws.merge_cells("A1:D1")
            ws["A1"].value = f"STATE: {state}"
            style_cell(ws["A1"], font=state_font, fill=state_fill, align=center_align)
            ws.row_dimensions[1].height = 20

            current_row = 2

            for district, centres_list in sorted(districts.items()):
                # District sub-header
                ws.merge_cells(f"A{current_row}:D{current_row}")
                ws[f"A{current_row}"].value = f"District: {district}  ({len(centres_list)} centres)"
                style_cell(
                    ws[f"A{current_row}"],
                    font=district_font,
                    fill=district_fill,
                    align=left_align,
                )
                ws.row_dimensions[current_row].height = 18
                current_row += 1

                # Column headers
                ws[f"A{current_row}"] = "AI Code"
                ws[f"B{current_row}"] = "Study Centre Name"
                ws[f"C{current_row}"] = "Address"
                ws[f"D{current_row}"] = "Valid"
                for col in ["A", "B", "C", "D"]:
                    style_cell(
                        ws[f"{col}{current_row}"],
                        font=Font(bold=True, size=9),
                        fill=PatternFill("solid", fgColor="AED6F1"),
                        align=center_align,
                        border=thin_border,
                    )
                current_row += 1

                # Data rows
                for centre in sorted(centres_list, key=lambda x: x.ai_code):
                    ws[f"A{current_row}"] = centre.ai_code  or ""
                    ws[f"B{current_row}"] = centre.name     or ""
                    ws[f"C{current_row}"] = centre.address  or ""
                    ws[f"D{current_row}"] = "✓" if centre.is_valid else "✗"
                    for col in ["A", "B", "C", "D"]:
                        style_cell(
                            ws[f"{col}{current_row}"],
                            align=left_align,
                            border=thin_border,
                        )
                    current_row += 1

                current_row += 1  # blank row between districts

            # Column widths
            ws.column_dimensions["A"].width = 12
            ws.column_dimensions["B"].width = 40
            ws.column_dimensions["C"].width = 50
            ws.column_dimensions["D"].width = 8

        wb.save(str(XLSX_PATH))

    # ── PDF Report ────────────────────────────────────────────

    @staticmethod
    def _pdf_safe(text: str) -> str:
        """
        Sanitise text for Helvetica (Latin-1 only) PDF rendering.
        Replaces common Unicode characters with ASCII equivalents.
        """
        if not text:
            return ""
        replacements = {
            '\u2013': '-',   # en-dash
            '\u2014': '-',   # em-dash
            '\u2018': "'",   # left single quote
            '\u2019': "'",   # right single quote
            '\u201c': '"',   # left double quote
            '\u201d': '"',   # right double quote
            '\u2026': '...', # ellipsis
            '\u00a0': ' ',   # non-breaking space
            '\u2022': '*',   # bullet
            '\u2010': '-',   # hyphen
            '\u2011': '-',   # non-breaking hyphen
            '\u2012': '-',   # figure dash
            '\u00b7': '.',   # middle dot
        }
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)
        # Final fallback: encode to Latin-1 with replace for any remaining
        return text.encode('latin-1', errors='replace').decode('latin-1')

    def _generate_pdf(self, grouped: Dict, all_centres: List) -> None:
        """Generate a structured PDF directory using fpdf2."""
        from fpdf import FPDF

        safe = self._pdf_safe  # shorthand

        class NIOSPdf(FPDF):
            def header(self):
                self.set_font("Helvetica", "B", 9)
                self.set_text_color(80, 80, 80)
                self.cell(0, 6, "NIOS Academic Study Centre Directory - India", align="C", new_x="LMARGIN", new_y="NEXT")
                self.set_draw_color(200, 200, 200)
                self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
                self.ln(2)

            def footer(self):
                self.set_y(-12)
                self.set_font("Helvetica", "I", 8)
                self.set_text_color(150, 150, 150)
                self.cell(0, 5, f"Page {self.page_no()} | Generated {datetime.now().strftime('%d %b %Y')}", align="C")

        pdf = NIOSPdf(orientation="P", unit="mm", format="A4")
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_margins(15, 15, 15)

        # ── Cover page ──
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 22)
        pdf.set_text_color(27, 79, 114)
        pdf.ln(20)
        pdf.cell(0, 12, "NIOS Academic Study Centre", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 12, "Directory - India", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(8)
        pdf.set_font("Helvetica", "", 12)
        pdf.set_text_color(80, 80, 80)
        pdf.cell(0, 8, f"Total Study Centres: {len(all_centres)}", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 8, f"States / UTs: {len(grouped)}", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 8, f"Category: Academic | Source: sdmis.nios.ac.in", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 8, f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}", align="C", new_x="LMARGIN", new_y="NEXT")

        # ── Data pages ──
        for state, districts in sorted(grouped.items()):
            pdf.add_page()

            # State heading
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_fill_color(27, 79, 114)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 9, safe(f"  {state.upper()}"), fill=True, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)

            for district, centres_list in sorted(districts.items()):
                # District sub-heading
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_fill_color(214, 234, 248)
                pdf.set_text_color(21, 67, 96)
                pdf.cell(0, 7, safe(f"  {district}  ({len(centres_list)} centres)"), fill=True, new_x="LMARGIN", new_y="NEXT")
                pdf.ln(1)

                # Table header
                pdf.set_font("Helvetica", "B", 8)
                pdf.set_fill_color(174, 214, 241)
                pdf.set_text_color(0, 0, 0)
                pdf.set_draw_color(180, 180, 180)
                pdf.cell(22,  6, "AI Code",  border=1, fill=True)
                pdf.cell(72,  6, "Study Centre Name", border=1, fill=True)
                pdf.cell(86,  6, "Address",  border=1, fill=True, new_x="LMARGIN", new_y="NEXT")

                # Data rows
                pdf.set_font("Helvetica", "", 7)
                pdf.set_fill_color(255, 255, 255)
                pdf.set_text_color(30, 30, 30)
                fill_alt = False
                for centre in sorted(centres_list, key=lambda x: x.ai_code):
                    if fill_alt:
                        pdf.set_fill_color(245, 250, 255)
                    else:
                        pdf.set_fill_color(255, 255, 255)
                    fill_alt = not fill_alt

                    # Multi-line cells using multi_cell with a fixed height approach
                    x = pdf.get_x()
                    y = pdf.get_y()

                    name_str    = safe((centre.name    or "")[:80])
                    address_str = safe((centre.address or "")[:120])
                    ai_code_str = safe((centre.ai_code or ""))

                    pdf.cell(22, 5, ai_code_str, border=1, fill=True)
                    pdf.cell(72, 5, name_str,    border=1, fill=True)
                    pdf.cell(86, 5, address_str, border=1, fill=True, new_x="LMARGIN", new_y="NEXT")

                pdf.ln(2)

            pdf.ln(3)

        pdf.output(str(PDF_PATH))

    # ── AI Executive Summary ───────────────────────────────────

    def generate_summary(self, total_centres: int, total_states: int, top_states: List) -> str:
        """
        Use Gemma 4 31B Free to generate a short executive summary
        for inclusion in the report cover page.
        """
        if not OPENROUTER_API_KEY:
            return ""
        try:
            client = _get_ai_client()
            prompt = (
                f"Write a 3-sentence executive summary for a NIOS (National Institute of Open Schooling) "
                f"Academic Study Centre Directory covering India. "
                f"Total centres: {total_centres} across {total_states} states/UTs. "
                f"Top states by count: {', '.join(top_states[:5])}. "
                f"Keep it professional and concise."
            )
            response = client.chat.completions.create(
                model=MODELS["report"],
                messages=[{"role": "user", "content": prompt}],
                max_tokens=200,
                temperature=0.7,
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            log.debug(f"Summary AI generation failed: {e}")
            return ""
